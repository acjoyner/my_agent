import json
import os
import re
from datetime import datetime
from github import Github
from docx import Document
from pypdf import PdfReader, PdfWriter
from tools.client import client
from config.settings import RESUME_TEXT

# --- Standard Skill Interface Requirements ---

def get_tools():
    return [
        {
            "name": "list_github_projects",
            "description": "Fetches a list of public repositories and descriptions from a GitHub user.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "username": {"type": "string", "description": "The GitHub username to fetch projects from."}
                },
                "required": ["username"]
            }
        },
        {
            "name": "propose_resume_update",
            "description": "Generates professional bullet points for selected projects and proposes them for the resume.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "projects": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "description": {"type": "string"},
                                "tech_stack": {"type": "string"}
                            }
                        }
                    }
                },
                "required": ["projects"]
            }
        },
        {
            "name": "approve_resume_update",
            "description": "Finalizes the resume update by writing approved bullets to the Word and PDF files and syncing with the Web App.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "approved_bullets": {"type": "string", "description": "The final text to insert into the resume."},
                    "projects_data": {
                        "type": "array",
                        "description": "The raw project data used for Web App synchronization.",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "description": {"type": "string"},
                                "tech_stack": {"type": "string"}
                            }
                        }
                    }
                },
                "required": ["approved_bullets"]
            }
        }
    ]

def handle_tool(name, inputs, memory):
    if name == "list_github_projects":
        return fetch_github_projects(inputs["username"])
    if name == "propose_resume_update":
        return propose_update(inputs["projects"])
    if name == "approve_resume_update":
        # Pass the raw projects data for web sync if available
        projects = inputs.get("projects_data", [])
        return finalize_update(inputs["approved_bullets"], projects)
    return None

# --- Core Logic ---

def fetch_github_projects(username):
    """Fetches public repo info and returns a readable summary."""
    try:
        g = Github()
        user = g.get_user(username)
        repos = []
        summary = f"### GitHub Projects for {username}\n\n"
        # Get top 5 most recently updated public repos
        for repo in user.get_repos(sort="updated", direction="desc")[:5]:
            repos.append({
                "name": repo.name,
                "description": repo.description or "No description",
                "tech_stack": repo.language or "Various"
            })
            summary += f"- **{repo.name}**: {repo.description or 'No description'} (Tech: {repo.language or 'Various'})\n"
        
        return json.dumps({
            "projects": repos,
            "display_text": summary + "\n\n**Would you like me to propose resume bullet points for these projects?**"
        })
    except Exception as e:
        return f"Error fetching GitHub data: {e}"

def propose_update(projects):
    """Generates professional bullet points using local models."""
    projects_str = ""
    for p in projects:
        name = p.get("name", "Unknown Project")
        desc = p.get("description") or p.get("focus") or "No description provided."
        tech = p.get("tech_stack") or "Various technologies."
        projects_str += f"- Project: {name}\n  Description: {desc}\n  Tech: {tech}\n\n"
    
    prompt = f"""
    You are an expert Executive Resume Writer.
    
    Task: Convert these projects into high-impact professional bullet points for a Senior Software Engineer.
    
    Guidelines:
    - Use strong action verbs (e.g., 'Architected', 'Spearheaded', 'Optimized').
    - Focus on measurable impact and technical complexity.
    
    Projects:
    {projects_str}
    
    Return the proposal in a clear markdown list.
    """
    
    response = client.messages.create(
        model="high-fidelity", 
        messages=[{"role": "user", "content": prompt}]
    )
    
    try:
        proposal = response.content[0].text
        return json.dumps({
            "status": "Awaiting Approval",
            "proposal": proposal,
            "display_text": f"### Proposed Resume Updates\n\n{proposal}\n\n**Do you approve these changes?** (Say 'Approve' to save).",
            "projects_data": projects
        })
    except Exception as e:
        return f"Error generating proposal: {e}"

def update_web_resume(new_projects):
    """Updates the Projects.tsx file with smart merging."""
    web_app_path = "/Users/anthonyjoyner/Documents/Projects/resume/resume-app/client/src/pages/Projects.tsx"
    if not os.path.exists(web_app_path):
        return f"Error: Could not find resume-app project at {web_app_path}"
    
    try:
        with open(web_app_path, 'r') as f:
            content = f.read()
            
        # 1. IDENTIFY EXISTING PROJECTS
        pattern = r"const projects = (\[.*?\]);"
        match = re.search(pattern, content, flags=re.DOTALL)
        
        existing_projects = []
        if match:
            js_array = match.group(1)
            # Try to parse JS as JSON
            js_json = re.sub(r'(\s)(\w+):', r'\1"\2":', js_array)
            js_json = re.sub(r",\s*\]", "]", js_json)
            js_json = re.sub(r",\s*\}", "}", js_json)
            try:
                existing_projects = json.loads(js_json)
            except:
                existing_projects = []

        # 2. SMART MERGE (Preserve by name, keep non-github items)
        project_map = {p.get("title"): p for p in existing_projects}
        
        for p in new_projects:
            name = p.get("name")
            desc = p.get("description") or p.get("focus") or "Project development."
            tech = p.get("tech_stack", "Various")
            
            if name in project_map:
                # Update existing entry but keep original period if possible
                project_map[name]["description"] = desc
                project_map[name]["skills"] = [tech]
            else:
                # Create new entry
                project_map[name] = {
                    "title": name,
                    "period": datetime.now().strftime("%Y") + " - Present",
                    "description": desc,
                    "skills": [tech]
                }
        
        final_list = list(project_map.values())
        
        # 3. WRITE BACK
        replacement = f"const projects = {json.dumps(final_list, indent=2)};"
        new_content = re.sub(pattern, replacement, content, flags=re.DOTALL)
        
        with open(web_app_path, 'w') as f:
            f.write(new_content)
            
        return f"Web App updated. Synced {len(new_projects)} projects into a total of {len(final_list)} entries."
    except Exception as e:
        return f"Error updating Web App: {e}"

def finalize_update(bullets, projects=[]):
    """Updates the .docx file and syncs to Web."""
    uploads_dir = "uploads"
    resume_path = None
    
    if os.path.exists(uploads_dir):
        files = [os.path.join(uploads_dir, f) for f in os.listdir(uploads_dir) 
                 if f.endswith('.docx') and "_Updated" not in f]
        if files:
            resume_path = max(files, key=os.path.getctime)
            
    if not resume_path:
        return "Error: Could not find a clean Word document in the uploads folder."
        
    results = []
    try:
        # 1. Update Word Doc
        doc = Document(resume_path)
        doc.add_heading('Flagship Projects (Synchronized via Agent)', level=2)
        
        clean_bullets = bullets.replace("### ", "").split("\n")
        for line in clean_bullets:
            line = line.strip()
            if not line: continue
            if line.startswith("-") or line.startswith("*"):
                try:
                    doc.add_paragraph(line[1:].strip(), style='List Bullet')
                except:
                    doc.add_paragraph(f"• {line[1:].strip()}")
            else:
                doc.add_paragraph(line)
        
        updated_path = resume_path.replace(".docx", "_Updated.docx")
        doc.save(updated_path)
        results.append(f"Word document updated.")
        
        # 2. Update Web App
        if projects:
            results.append(update_web_resume(projects))
            results.append("Manual deploy recommended: `cd /Users/anthonyjoyner/Documents/Projects/resume/resume-app && npm run build && firebase deploy`")
        
        return json.dumps({
            "status": "Success",
            "message": "Updates completed: " + " | ".join(results),
            "display_text": "### Update Complete\n\n" + "\n".join([f"- {r}" for r in results])
        })
    except Exception as e:
        return f"Error updating document: {e}"

def get_system_prompt():
    return "You have the 'Resume Auto-Updater' skill. You can sync GitHub projects and update both Word files and your Firebase Web App with user approval."
